import requests
import time
import json
import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

headers = {

}
url = 'https://mp.weixin.qq.com/cgi-bin/appmsgpublish'


def download_images(url, date):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'lxml')
    meta_element = soup.find('meta', {'property': 'og:title'})
    title = meta_element['content']
    title = re.sub('[/\:*?<>"|]', ' ', title)
    title = time.strftime('%Y.%m.%d', date) + " " + title
    os.mkdir(title)
    os.chdir(title)
    os.mkdir(title)
    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(10.5)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    doc.add_paragraph(title)
    tags = soup.find_all(lambda tag: tag.name == 'p' and
                                     tag.has_attr('style') and
                                     'text-wrap: wrap' in tag['style'] or tag.name == 'img')
    flag = 0
    index = 0
    for tag in tags:
        if tag.name == 'img':
            image_url = tag.get('data-src')
            if image_url is not None:
                if flag > 1 and image_url.endswith("jpeg"):
                    response = requests.get(image_url)
                    index = index + 1
                    with open(os.path.join(title, str(index) + ".jpg"), 'wb') as f:
                        f.write(response.content)
                    doc.add_picture(os.path.join(title, str(index) + ".jpg"), width=Inches(3))
            flag = flag + 1
        elif "文稿撰写：" in tag.text or "稿件来源：" in tag.text:
            break
        elif flag>0:
            content = tag.text
            doc.add_paragraph(content)
    # if not os.path.exists(title):
    doc.save(title + '.docx')
    os.chdir('..')


def page(num=1):
    title = []
    link = []
    sent_time = []
    for i in range(num):
        data = {
            'sub': 'list',
            'begin': 0,
            'count': 30,
            'query': '',
            'type': '101_1',
            'show_type': '',
            'free_publish_type': 1,
            'sub_action': 'list_ex',
            'search_card': 1,
            'token': 1061839601,
            'lang': 'zh_CN',
            'f': 'json',
            'ajax': 1
        }
        r = requests.get(url, headers=headers, params=data)
        dic = json.loads(r.json()['publish_page'])
        for page in dic['publish_list']:
            page_info = json.loads(page['publish_info'])
            if page_info['type'] == 9:
                page_time = time.localtime(page_info['sent_info']['time'])
                sent_time.append(page_time)
                title.append(page_info['appmsgex'][0]['title'])
                link.append(page_info['appmsgex'][0]['link'])
            if page_info['type'] == 1:
                page_time = time.localtime(page_info['publish_info']['update_time'])
                sent_time.append(page_time)
                title.append(page_info['appmsgex'][0]['title'])
                link.append(page_info['appmsgex'][0]['link'])
    return (sent_time, title, link)


if __name__ == '__main__':

    (tim, tle, lik) = page(5)
    for id, (x, y, z) in enumerate(zip(tim, tle, lik), start=1):
        print(id, time.strftime('%Y.%m.%d', x), y)
    id1 = int(input('Enter id1: '))
    id2 = int(input('Enter id2: '))
    for i in range(id1 - 1, id2):
        download_images(lik[i], x)
        # download_images("https://mp.weixin.qq.com/s/wYae_8I1o9kMnZBzdmixAA", x)
        print(tle[i] + ' has been downloaded!')
